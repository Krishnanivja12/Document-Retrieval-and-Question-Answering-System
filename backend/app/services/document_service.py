import uuid
from fastapi import UploadFile
from sqlalchemy.orm import Session
from loguru import logger
from app.models.document import Document
from app.utils.file_handler import save_upload, TEXT_DIR
from app.utils.url_parser import fetch_text_from_url
from app.core.errors import AppException

def extract_text_from_pdf(file_path: str) -> str:
    import fitz  # PyMuPDF
    doc = fitz.open(file_path)
    full_text = []
    for page in doc:
        full_text.append(page.get_text())
    doc.close()
    return "\n\n".join(full_text)

def extract_text_from_docx(file_path: str) -> str:
    from docx import Document as DocxDocument
    doc = DocxDocument(file_path)
    paragraphs = [para.text for para in doc.paragraphs]
    return "\n".join(paragraphs)

def extract_text(file_path: str, ext: str) -> str:
    if ext == ".pdf":
        return extract_text_from_pdf(file_path)
    elif ext == ".docx":
        return extract_text_from_docx(file_path)
    elif ext == ".txt":
        with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
            return f.read()
    else:
        raise AppException(detail=f"No extraction method for {ext}", status_code=400)

async def process_file(upload: UploadFile, db: Session) -> Document:
    file_info = await save_upload(upload)
    original_path = file_info["file_path"]
    # Extract text
    text = extract_text(original_path, file_info["ext"])
    TEXT_DIR.mkdir(parents=True, exist_ok=True)
    text_filename = f"{uuid.uuid4()}.txt"
    text_path = TEXT_DIR / text_filename
    with open(text_path, "w", encoding="utf-8") as f:
        f.write(text)
    doc = Document(
        filename=file_info["original_filename"],
        source_type="file",
        content_type=upload.content_type or "application/octet-stream",
        ext=file_info["ext"],
        text_path=str(text_path),
        size_bytes=file_info["size_bytes"],
        original_path=original_path,   # NEW
    )
    db.add(doc)
    db.commit()
    db.refresh(doc)
    # Chunk the document
    from app.services.chunking_service import chunk_document
    chunk_document(doc, db)
    return doc

async def process_url(url: str, db: Session) -> Document:
    # Fetch and clean text
    text = await fetch_text_from_url(url)
    # Save text
    TEXT_DIR.mkdir(parents=True, exist_ok=True)
    text_filename = f"{uuid.uuid4()}.txt"
    text_path = TEXT_DIR / text_filename
    with open(text_path, "w", encoding="utf-8") as f:
        f.write(text)
    # Determine an appropriate filename from URL
    filename = url.rstrip("/").split("/")[-1] or "webpage"
    ext = ".url.txt"  # pseudo extension
    doc = Document(
        filename=filename,
        source_type="url",
        source_url=url,
        content_type="text/html",
        original_path=None,
        ext=ext,
        text_path=str(text_path),
        size_bytes=len(text.encode("utf-8")),
    )
    db.add(doc)
    db.commit()
    db.refresh(doc)

    from app.services.chunking_service import chunk_document
    chunk_document(doc, db)

    logger.info(f"Ingested URL document {doc.id}: {url}")
    return doc
